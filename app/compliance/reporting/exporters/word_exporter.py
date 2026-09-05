"""Word 报告导出 — python-docx。

输入：ReviewDetailResponse 风格的 dict / Pydantic 对象
输出：.docx 文件路径
"""

from __future__ import annotations

import os
from datetime import datetime

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from app.config import settings
from app.logging_config import get_logger

logger = get_logger(__name__)

_LEVEL_COLOR = {
    "high": RGBColor(0xC0, 0x39, 0x2B),
    "medium": RGBColor(0xD6, 0x89, 0x10),
    "low": RGBColor(0x2E, 0x86, 0xAB),
}
_LEVEL_LABEL = {"high": "高风险", "medium": "中风险", "low": "低风险"}
_CATEGORY_LABEL = {
    "legality": "合法性",
    "equality": "对等性",
    "clarity": "明确性",
    "completeness": "完整性",
    "reasonableness": "合理性",
}


def export_word(review_data: dict, out_dir: str | None = None) -> str:
    """将审查结果导出为 Word，返回文件绝对路径。"""

    out_dir = out_dir or settings.compliance_report_dir
    os.makedirs(out_dir, exist_ok=True)

    review_id = review_data.get("review_id", "unknown")
    doc = Document()

    # ---- 全局样式 ----
    _setup_styles(doc)

    # ============== 封面 ==============
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("合规审查报告")
    run.bold = True
    run.font.size = Pt(28)

    doc.add_paragraph()
    doc.add_paragraph()

    _add_meta_row(doc, "审查任务 ID", review_id)
    _add_meta_row(doc, "合同类型", review_data.get("doc_type") or "—")
    _add_meta_row(
        doc,
        "完成时间",
        (review_data.get("completed_at") or datetime.utcnow().isoformat())[:19],
    )
    h = review_data.get("high_risk_count", 0)
    m = review_data.get("medium_risk_count", 0)
    l = review_data.get("low_risk_count", 0)
    _add_meta_row(doc, "风险统计", f"🔴 {h}  🟡 {m}  🟢 {l}  合计 {h + m + l}")

    doc.add_page_break()

    # ============== 执行摘要 ==============
    doc.add_heading("执行摘要", level=1)
    total = h + m + l
    if total == 0:
        p = doc.add_paragraph("本次审查未检出风险条款，合同合规。")
    else:
        p = doc.add_paragraph(
            f"本次审查共识别 **{total}** 项风险，其中高风险 {h} 项、中风险 {m} 项、低风险 {l} 项。"
            f"{'建议在签署前逐条处理。' if h > 0 else '整体风险可控。'}"
        )
    p.runs[0].font.size = Pt(11)

    summary = review_data.get("summary")
    if summary:
        doc.add_paragraph(summary)

    doc.add_page_break()

    # ============== 合同基本信息 ==============
    doc.add_heading("合同基本信息", level=1)
    key_info = review_data.get("key_info") or {}
    if key_info:
        table = doc.add_table(rows=0, cols=2)
        table.style = "Light Grid Accent 1"
        _KEY_LABELS = {
            "party_a": "甲方",
            "party_b": "乙方",
            "sign_date": "签订日期",
            "effective_date": "生效日期",
            "term": "合同期限",
            "amount": "合同金额",
            "payment_method": "付款方式",
            "penalty_cap": "违约金上限",
            "ip_ownership": "知识产权归属",
            "confidentiality_period": "保密期限",
            "dispute_resolution": "争议解决方式",
        }
        for k, v in key_info.items():
            if v:
                row = table.add_row().cells
                row[0].text = _KEY_LABELS.get(k, k)
                row[1].text = str(v)
    else:
        doc.add_paragraph("（未提取到结构化合同信息）")

    doc.add_page_break()

    # ============== 风险总览 ==============
    doc.add_heading("风险总览", level=1)
    risks = review_data.get("risks", [])
    if not risks:
        doc.add_paragraph("未检出任何风险条款。")
    else:
        _add_risks_table(doc, risks)

    doc.add_page_break()

    # ============== 逐条风险详情 ==============
    doc.add_heading("逐条风险详情", level=1)

    # 按风险等级排序：high → medium → low
    _ORDER = {"high": 0, "medium": 1, "low": 2}
    risks_sorted = sorted(risks, key=lambda r: _ORDER.get(r.get("risk_level", "low"), 9))

    for idx, risk in enumerate(risks_sorted, 1):
        level = risk.get("risk_level", "low")
        level_label = _LEVEL_LABEL.get(level, level)
        color = _LEVEL_COLOR.get(level, RGBColor(0, 0, 0))

        h2 = doc.add_heading(
            f"#{idx} [{level_label}] {risk.get('clause_number') or '条款未定位'}",
            level=2,
        )
        for run in h2.runs:
            run.font.color.rgb = color

        cat = _CATEGORY_LABEL.get(risk.get("risk_category", ""), risk.get("risk_category", "—"))
        doc.add_paragraph(f"风险类别：{cat}    置信度：{risk.get('ai_confidence', '—')}")

        doc.add_paragraph("风险描述：").runs[0].bold = True
        doc.add_paragraph(risk.get("description", "—"))

        if risk.get("suggestion"):
            p = doc.add_paragraph("修改建议：").runs[0].bold = True
            doc.add_paragraph(risk["suggestion"])
        if risk.get("suggestion_reason"):
            doc.add_paragraph(f"修改理由：{risk['suggestion_reason']}")

        refs = risk.get("legal_references", [])
        if refs:
            doc.add_paragraph("法规依据：").runs[0].bold = True
            for ref in refs:
                verified = "✅" if ref.get("verified") else "⚠️ 待核实"
                art = f" {ref['ref_article']}" if ref.get("ref_article") else ""
                doc.add_paragraph(f"{verified} {ref['ref_name']}{art}", style="List Bullet")
                if ref.get("ref_content"):
                    doc.add_paragraph(f"    {ref['ref_content'][:200]}", style="List Bullet 2")

    doc.add_page_break()

    # ============== 法规附录 ==============
    doc.add_heading("法规附录", level=1)
    _render_appendix(doc, risks)

    # ---- 保存 ----
    fname = f"compliance_report_{review_id[:8]}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
    fpath = os.path.join(out_dir, fname)
    doc.save(fpath)
    logger.info("Word report saved: %s", fpath)
    return fpath


def _setup_styles(doc: Document):
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)


def _add_meta_row(doc: Document, label: str, value: str):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}：")
    r1.bold = True
    p.add_run(value)


def _add_risks_table(doc: Document, risks: list[dict]):
    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, t in enumerate(["#", "风险等级", "风险类别", "条款", "描述"]):
        hdr[i].text = t
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True

    for idx, risk in enumerate(risks, 1):
        cells = table.add_row().cells
        level = risk.get("risk_level", "low")
        cells[0].text = str(idx)
        cells[1].text = _LEVEL_LABEL.get(level, level)
        for r in cells[1].paragraphs[0].runs:
            r.font.color.rgb = _LEVEL_COLOR.get(level, RGBColor(0, 0, 0))
        cells[2].text = _CATEGORY_LABEL.get(
            risk.get("risk_category", ""), risk.get("risk_category", "—")
        )
        cells[3].text = risk.get("clause_number") or "—"
        cells[4].text = (risk.get("description") or "")[:80]


def _render_appendix(doc: Document, risks: list[dict]):
    seen: dict[str, dict] = {}
    for risk in risks:
        for ref in risk.get("legal_references", []):
            key = f"{ref.get('ref_name')}|{ref.get('ref_article', '')}"
            if key not in seen:
                seen[key] = ref

    if not seen:
        doc.add_paragraph("本次审查未引用法规条文。")
        return

    for key, ref in seen.items():
        name = ref.get("ref_name", "未知法规")
        art = ref.get("ref_article")
        heading = f"{name}" + (f"  {art}" if art else "")
        doc.add_heading(heading, level=2)
        doc.add_paragraph(ref.get("ref_content", "—"))
